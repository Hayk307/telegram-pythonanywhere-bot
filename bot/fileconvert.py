"""File conversion helpers for the /convertfile command.

Heavy third-party libraries (Pillow, pypdf, python-docx, fpdf, ffmpeg) are
imported *lazily inside* the conversion functions, never at module import
time. That way the bot still boots and every other command keeps working
even when these optional dependencies aren't installed on the host yet — a
conversion attempt just fails with a clear message instead of taking the
whole worker down.

Conversions only happen *within* a family:
  - image  <-> image  (Pillow)
  - media  <-> media   (ffmpeg; audio and video share one family, so
                        video->audio like mp4->mp3 is just a media->media)
  - doc    <-> doc     (text-only: extract text, then re-emit it)

Cross-family requests (e.g. pdf->mp3) are impossible and rejected.
"""

from __future__ import annotations

import os
import re
import subprocess

IMAGE_EXTS = {"png", "jpg", "jpeg", "webp", "bmp", "gif", "tiff"}
AUDIO_EXTS = {"mp3", "wav", "m4a", "aac", "ogg", "flac"}
VIDEO_EXTS = {"mp4", "mov", "webm", "mkv", "avi"}
MEDIA_EXTS = AUDIO_EXTS | VIDEO_EXTS
# Document formats we can *read* (sources) and *write* (targets). We can read
# .md but only emit pdf/docx/txt.
DOC_SOURCE_EXTS = {"pdf", "docx", "txt", "md"}
DOC_TARGET_EXTS = {"pdf", "docx", "txt"}

SUPPORTED_TARGETS = IMAGE_EXTS | MEDIA_EXTS | DOC_TARGET_EXTS

# ffmpeg gets a hard timeout so a pathological file can't wedge the worker.
FFMPEG_TIMEOUT = 120


class ConversionError(Exception):
    """Raised with a user-facing message when a file can't be converted."""


def normalize_ext(name_or_ext: str) -> str:
    """Return a bare lowercase extension from a filename, ext, or '.ext'."""
    ext = (name_or_ext or "").lower().strip()
    if "." in ext:
        ext = ext.rsplit(".", 1)[-1]
    return ext.lstrip(".")


def category_of(ext: str) -> str | None:
    """Map an extension to its conversion family, or None if unsupported."""
    ext = normalize_ext(ext)
    if ext in IMAGE_EXTS:
        return "image"
    if ext in MEDIA_EXTS:
        return "media"
    if ext in DOC_SOURCE_EXTS:
        return "doc"
    return None


def parse_target_format(caption: str) -> str | None:
    """Pull the desired output extension out of a file's caption.

    Accepts "mp3", ".mp3", "to mp3", "convert to png", "/convertfile pdf",
    or a short phrase — it prefers a token that is a known target format and
    otherwise falls back to the last word (so the caller can reject it with a
    helpful message).
    """
    if not caption:
        return None
    text = caption.strip().lower()
    text = re.sub(r"^/convertfile(@\w+)?\b", "", text)
    parts = [p for p in re.split(r"[\s.]+", text) if p and p not in ("to", "convert", "into", "as", "a", "an")]
    if not parts:
        return None
    known = [p for p in parts if p in SUPPORTED_TARGETS]
    return known[-1] if known else parts[-1]


def convert(input_path: str, source_ext: str, target_ext: str) -> str:
    """Convert input_path into target_ext, returning the output file path.

    Raises ConversionError (with a user-friendly message) on any problem.
    """
    source_ext = normalize_ext(source_ext)
    target_ext = normalize_ext(target_ext)
    src_cat = category_of(source_ext)
    if target_ext not in SUPPORTED_TARGETS:
        raise ConversionError(f"I can't produce .{target_ext} files.")
    if src_cat is None:
        raise ConversionError(f"I can't read .{source_ext or '?'} files.")
    tgt_cat = category_of(target_ext)
    if tgt_cat != src_cat:
        raise ConversionError(
            f"Can't turn a {src_cat} file (.{source_ext}) into a {tgt_cat} file "
            f"(.{target_ext}). I convert within a family: image↔image, "
            "audio/video↔audio/video, document↔document."
        )
    if source_ext == target_ext:
        raise ConversionError(f"That file is already a .{target_ext}.")

    output_path = f"{os.path.splitext(input_path)[0]}.{target_ext}"
    if src_cat == "image":
        _convert_image(input_path, output_path, target_ext)
    elif src_cat == "media":
        _convert_media(input_path, output_path)
    else:
        _convert_doc(input_path, source_ext, output_path, target_ext)
    if not os.path.exists(output_path):
        raise ConversionError("Conversion produced no output.")
    return output_path


def _convert_image(input_path: str, output_path: str, target_ext: str) -> None:
    try:
        from PIL import Image
    except ImportError:
        raise ConversionError("Image conversion isn't available (missing Pillow).")
    with Image.open(input_path) as img:
        fmt = "JPEG" if target_ext in ("jpg", "jpeg") else target_ext.upper()
        # JPEG and BMP have no alpha channel — flatten to RGB first.
        if fmt in ("JPEG", "BMP") and img.mode in ("RGBA", "LA", "P"):
            img = img.convert("RGB")
        img.save(output_path, fmt)


def _ffmpeg_exe() -> str:
    """Return a path to an ffmpeg binary: the pip-bundled one if present,
    otherwise plain 'ffmpeg' from PATH."""
    try:
        import imageio_ffmpeg

        return imageio_ffmpeg.get_ffmpeg_exe()
    except Exception:
        return "ffmpeg"


def _convert_media(input_path: str, output_path: str) -> None:
    exe = _ffmpeg_exe()
    try:
        proc = subprocess.run(
            [exe, "-y", "-i", input_path, output_path],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=FFMPEG_TIMEOUT,
        )
    except FileNotFoundError:
        raise ConversionError(
            "Audio/video conversion isn't available (ffmpeg not found)."
        )
    except subprocess.TimeoutExpired:
        raise ConversionError("That media file took too long to convert.")
    if proc.returncode != 0:
        raise ConversionError("ffmpeg couldn't convert that file.")


def _extract_text(input_path: str, source_ext: str) -> str:
    if source_ext == "pdf":
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ConversionError("PDF reading isn't available (missing pypdf).")
        reader = PdfReader(input_path)
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if source_ext == "docx":
        try:
            import docx
        except ImportError:
            raise ConversionError(
                "Word reading isn't available (missing python-docx)."
            )
        document = docx.Document(input_path)
        return "\n".join(p.text for p in document.paragraphs)
    # txt / md — plain text.
    with open(input_path, "r", encoding="utf-8", errors="replace") as handle:
        return handle.read()


def _convert_doc(
    input_path: str, source_ext: str, output_path: str, target_ext: str
) -> None:
    text = _extract_text(input_path, source_ext)
    lines = text.splitlines() or [""]
    if target_ext == "txt":
        with open(output_path, "w", encoding="utf-8") as handle:
            handle.write(text)
    elif target_ext == "docx":
        try:
            import docx
        except ImportError:
            raise ConversionError(
                "Word writing isn't available (missing python-docx)."
            )
        document = docx.Document()
        for line in lines:
            document.add_paragraph(line)
        document.save(output_path)
    elif target_ext == "pdf":
        try:
            from fpdf import FPDF
        except ImportError:
            raise ConversionError("PDF writing isn't available (missing fpdf2).")
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        for line in lines:
            # fpdf2's built-in fonts are Latin-1 only, so non-Latin text (e.g.
            # Armenian) is replaced rather than crashing. Bundling a Unicode
            # TTF via add_font() would lift this limit.
            safe = line.encode("latin-1", "replace").decode("latin-1")
            # w=epw + new_x=LMARGIN keeps each line full-width and returns the
            # cursor to the left margin (the default new_x=RIGHT would leave it
            # at the right margin, so the next line would have zero width).
            pdf.multi_cell(w=pdf.epw, h=8, text=safe, new_x="LMARGIN", new_y="NEXT")
        pdf.output(output_path)
